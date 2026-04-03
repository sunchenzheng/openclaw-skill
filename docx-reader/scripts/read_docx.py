#!/usr/bin/env python3
"""
docx_reader.py — 读取 Word 文档内容
用法：
  python read_docx.py <docx路径> [--extract-images] [--output-json]
"""

import sys
import os
import json
import base64
import argparse
from pathlib import Path

def read_docx(docx_path, extract_images=False, output_json=False):
    """读取 .docx 文件，返回结构化内容"""
    try:
        from docx import Document
    except ImportError:
        print("Error: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document(docx_path)
    base_dir = Path(docx_path).parent
    stem = Path(docx_path).stem

    result = {
        "file": str(docx_path),
        "paragraphs": [],
        "tables": [],
        "images": [],
        "headings": []
    }

    # 读取段落
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name
        result["paragraphs"].append({
            "index": i,
            "style": style,
            "text": text
        })
        if style.startswith("Heading") or style == "Title":
            result["headings"].append({
                "style": style,
                "text": text
            })

    # 读取表格
    for t_idx, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(row_data)
        result["tables"].append({
            "index": t_idx,
            "rows": len(table_data),
            "cols": len(table_data[0]) if table_data else 0,
            "data": table_data
        })

    # 提取图片
    if extract_images:
        img_dir = base_dir / f"{stem}_images"
        img_dir.mkdir(exist_ok=True)

        for i, rel in enumerate(doc.part.rels.values()):
            if "image" in rel.target_ref:
                img_data = rel.target_part.blob
                ext = rel.target_ref.split(".")[-1].lower()
                if ext not in ["png", "jpg", "jpeg", "gif", "bmp", "webp"]:
                    ext = "png"
                img_name = f"image_{i+1}.{ext}"
                img_path = img_dir / img_name
                with open(img_path, "wb") as f:
                    f.write(img_data)
                result["images"].append({
                    "index": i + 1,
                    "filename": img_name,
                    "path": str(img_path),
                    "size_bytes": len(img_data)
                })

    return result


def print_summary(result):
    """打印内容摘要"""
    print(f"\n{'='*50}")
    print(f"📄 文件：{result['file']}")
    print(f"{'='*50}")
    print(f"\n📌 标题/段落结构：")
    for h in result["headings"]:
        indent = "  " if h["style"] != "Title" else ""
        print(f"  {indent}[{h['style']}] {h['text']}")

    print(f"\n📊 内容统计：")
    print(f"  段落数：{len(result['paragraphs'])}")
    print(f"  表格数：{len(result['tables'])}")
    print(f"  图片数：{len(result['images'])}")

    if result["tables"]:
        print(f"\n📋 表格内容（前3个）：")
        for t in result["tables"][:3]:
            print(f"  表格{t['index']+1}：{t['rows']}行 × {t['cols']}列")

    if result["images"]:
        print(f"\n🖼️ 图片列表：")
        for img in result["images"]:
            print(f"  {img['filename']} → {img['path']} ({img['size_bytes']} bytes)")

    print(f"\n{'='*50}")
    print("📝 全部段落内容：")
    print("="*50)
    for p in result["paragraphs"]:
        print(f"[{p['style']}] {p['text']}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="读取 Word 文档内容")
    parser.add_argument("docx_path", help=".docx 文件路径")
    parser.add_argument("--extract-images", action="store_true", help="同时提取文档中的图片")
    parser.add_argument("--json", action="store_true", help="输出 JSON 格式")
    args = parser.parse_args()

    if not os.path.exists(args.docx_path):
        print(f"Error: File not found: {args.docx_path}", file=sys.stderr)
        sys.exit(1)

    result = read_docx(args.docx_path, extract_images=args.extract_images)

    if args.json:
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        print_summary(result)
